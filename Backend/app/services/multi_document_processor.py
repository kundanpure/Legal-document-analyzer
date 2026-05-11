"""
Multi-format document processor with intelligent chunking
Enhanced with fallbacks and comprehensive error handling
"""

import asyncio
import hashlib
import uuid
import mimetypes
import re
from dotenv import load_dotenv
load_dotenv() 
from typing import Dict, List, Any, Optional, Tuple
from datetime import datetime, timezone
from pathlib import Path
from io import BytesIO

# Create a simple logger if config.logging is not available
try:
    from config.logging import get_logger
    logger = get_logger(__name__)
except ImportError:
    import logging
    logger = logging.getLogger(__name__)
    logger.setLevel(logging.INFO)

# Try to import optional dependencies with fallbacks
try:
    import aiofiles
    AIOFILES_AVAILABLE = True
except ImportError:
    AIOFILES_AVAILABLE = False
    logger.warning("aiofiles not available - using synchronous file operations")

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not available - DOCX processing will be limited")

try:
    from PyPDF2 import PdfReader
    PYPDF2_AVAILABLE = True
except ImportError:
    try:
        import PyPDF2
        PdfReader = PyPDF2.PdfFileReader
        PYPDF2_AVAILABLE = True
    except ImportError:
        PYPDF2_AVAILABLE = False
        logger.warning("PyPDF2 not available - PDF processing will be limited")

try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False
    logger.warning("pdfplumber not available - using fallback PDF processing")

try:
    import magic
    MAGIC_AVAILABLE = True
except ImportError:
    MAGIC_AVAILABLE = False
    logger.warning("python-magic not available - using filename-based type detection")

# Try to import custom modules with fallbacks
try:
    from app.services.intelligent_chunker import IntelligentChunker
    from app.services.intelligent_chunker import chunk_document_simple
    CHUNKER_AVAILABLE = True
except ImportError:
    CHUNKER_AVAILABLE = False
    logger.warning("IntelligentChunker not available - using basic chunking")

try:
    from app.models.multi_document_models import StandardDocument, DocumentChunk
    MODELS_AVAILABLE = True
except ImportError:
    MODELS_AVAILABLE = False
    logger.warning("Multi-document models not available - using fallback models")

# Custom exceptions
class DocumentProcessingError(Exception):
    """Exception raised for document processing errors"""
    pass

# Mock models when real models aren't available
class MockStandardDocument:
    def __init__(self, **kwargs):
        self.id = kwargs.get('id', f"doc_{uuid.uuid4().hex[:12]}")
        self.session_id = kwargs.get('session_id', 'default_session')
        self.filename = kwargs.get('filename', 'document.txt')
        self.content_type = kwargs.get('content_type', 'text/plain')
        self.content = kwargs.get('content', '')
        self.metadata = kwargs.get('metadata', {})
        self.page_count = kwargs.get('page_count', 1)
        self.word_count = kwargs.get('word_count', 0)
        self.file_hash = kwargs.get('file_hash', '')
        self.created_at = kwargs.get('created_at', datetime.now(timezone.utc))
        
        # Calculate word count if not provided
        if not self.word_count and self.content:
            self.word_count = len(self.content.split())
    
    def to_dict(self):
        return {
            'id': self.id,
            'filename': self.filename,
            'content_type': self.content_type,
            'page_count': self.page_count,
            'word_count': self.word_count
        }

class MockDocumentChunk:
    def __init__(self, **kwargs):
        self.id = kwargs.get('id', f"chunk_{uuid.uuid4().hex[:8]}")
        self.document_id = kwargs.get('document_id', '')
        self.chunk_index = kwargs.get('chunk_index', 0)
        self.total_chunks = kwargs.get('total_chunks', 1)
        self.content = kwargs.get('content', '')
        self.page_range = kwargs.get('page_range', '1-1')
        self.char_range = kwargs.get('char_range', '0-0')
        self.word_count = kwargs.get('word_count', 0)
        self.metadata = kwargs.get('metadata', {})
        self.created_at = kwargs.get('created_at', datetime.now(timezone.utc))
        
        if not self.word_count and self.content:
            self.word_count = len(self.content.split())
    
    def to_dict(self):
        return {
            'id': self.id,
            'document_id': self.document_id,
            'content': self.content[:100] + '...' if len(self.content) > 100 else self.content,
            'word_count': self.word_count
        }

# Use real models if available, otherwise use mocks
StandardDocument = StandardDocument if MODELS_AVAILABLE else MockStandardDocument
DocumentChunk = DocumentChunk if MODELS_AVAILABLE else MockDocumentChunk

class MultiDocumentProcessor:
    """
    Universal document processor supporting PDF, DOCX, TXT, DOC with fallbacks
    """
    
    def __init__(self):
        self.logger = logger
        
        # Initialize chunker with fallback
        if CHUNKER_AVAILABLE:
            try:
                self.chunker = IntelligentChunker(max_pages_per_chunk=15)
            except Exception as e:
                self.logger.warning(f"Failed to initialize IntelligentChunker: {e}")
                self.chunker = None
        else:
            self.chunker = None
            
        # Define supported formats with availability
        self.supported_formats = {
            'application/pdf': self.process_pdf if PYPDF2_AVAILABLE or PDFPLUMBER_AVAILABLE else self.process_pdf_fallback,
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': self.process_docx if DOCX_AVAILABLE else self.process_docx_fallback,
            'application/msword': self.process_doc,
            'text/plain': self.process_txt,
            'text/rtf': self.process_rtf,
            'text/markdown': self.process_txt,
            'application/octet-stream': self.process_unknown
        }
        
        pdf_status = '✓' if PYPDF2_AVAILABLE or PDFPLUMBER_AVAILABLE else '✗'
        docx_status = '✓' if DOCX_AVAILABLE else '✗'
        self.logger.info(f"MultiDocumentProcessor initialized - PDF: {pdf_status}, DOCX: {docx_status}")
        
    async def process_document(self, file_content: bytes, filename: str, session_id: str) -> Tuple[MockStandardDocument, List[MockDocumentChunk]]:
        """
        Process any supported document format
        
        Args:
            file_content: Document content as bytes
            filename: Original filename
            session_id: Session identifier
            
        Returns:
            Tuple of (StandardDocument, List[DocumentChunk])
        """
        
        try:
            self.logger.info(f"Processing document: {filename} for session: {session_id}")
            
            # Detect document type
            mime_type = self._detect_mime_type(file_content, filename)
            self.logger.info(f"Detected MIME type: {mime_type} for {filename}")
            
            # Get processor for this type
            processor = self.supported_formats.get(mime_type)
            if not processor:
                # Try fallback based on file extension
                processor = self._get_processor_by_extension(filename)
            
            if not processor:
                raise DocumentProcessingError(f"Unsupported document type: {mime_type}")
            
            # Extract text and metadata
            extracted_data = await processor(file_content, filename)
            
            # Create standardized document object
            document = StandardDocument(
                id=self._generate_document_id(),
                session_id=session_id,
                filename=filename,
                content_type=mime_type,
                content=extracted_data['text'],
                metadata=extracted_data.get('metadata', {}),
                page_count=extracted_data.get('page_count', 1),
                word_count=len(extracted_data['text'].split()) if extracted_data['text'] else 0,
                file_hash=self._calculate_file_hash(file_content),
                created_at=datetime.now(timezone.utc)
            )
            
            # Apply intelligent chunking
            chunks = await self._create_chunks(document)
            
            self.logger.info(f"Document processed: {document.id}, created {len(chunks)} chunks")
            
            return document, chunks
            
        except Exception as e:
            self.logger.error(f"Error processing document {filename}: {str(e)}")
            raise DocumentProcessingError(f"Failed to process document: {str(e)}")
    
    async def _create_chunks(self, document) -> List[MockDocumentChunk]:
        """Create chunks using available chunker or fallback"""
        
        if self.chunker and CHUNKER_AVAILABLE:
            try:
                return await self.chunker.create_chunks(document)
            except Exception as e:
                self.logger.error(f"Intelligent chunking failed: {e}, using fallback")
        
        # Fallback chunking
        return self._create_basic_chunks(document)
    
    def _create_basic_chunks(self, document) -> List[MockDocumentChunk]:
        """Basic chunking fallback"""
        
        max_chunk_size = 8000
        content = document.content
        
        if len(content) <= max_chunk_size:
            return [DocumentChunk(
                id=f"{document.id}_chunk_0",
                document_id=document.id,
                chunk_index=0,
                total_chunks=1,
                content=content,
                page_range="1-" + str(document.page_count),
                char_range=f"0-{len(content)}",
                word_count=len(content.split()),
                metadata={'chunk_type': 'complete'},
                created_at=datetime.now(timezone.utc)
            )]
        
        # Split into multiple chunks
        chunks = []
        chunk_count = (len(content) + max_chunk_size - 1) // max_chunk_size
        
        for i in range(chunk_count):
            start_pos = i * max_chunk_size
            end_pos = min((i + 1) * max_chunk_size, len(content))
            chunk_content = content[start_pos:end_pos]
            
            chunk = DocumentChunk(
                id=f"{document.id}_chunk_{i}",
                document_id=document.id,
                chunk_index=i,
                total_chunks=chunk_count,
                content=chunk_content,
                page_range=f"~{i+1}",
                char_range=f"{start_pos}-{end_pos}",
                word_count=len(chunk_content.split()),
                metadata={'chunk_type': 'basic'},
                created_at=datetime.now(timezone.utc)
            )
            
            chunks.append(chunk)
        
        return chunks
    
    async def process_pdf(self, content: bytes, filename: str) -> Dict[str, Any]:
        """Process PDF document with multiple libraries"""
        
        # Try pdfplumber first (better for complex PDFs)
        if PDFPLUMBER_AVAILABLE:
            try:
                return await self._process_pdf_with_pdfplumber(content, filename)
            except Exception as e:
                self.logger.warning(f"pdfplumber failed: {e}, trying PyPDF2")
        
        # Try PyPDF2 as fallback
        if PYPDF2_AVAILABLE:
            try:
                return await self._process_pdf_with_pypdf2(content, filename)
            except Exception as e:
                self.logger.warning(f"PyPDF2 failed: {e}, using basic extraction")
        
        # Final fallback
        return await self.process_pdf_fallback(content, filename)
    
    async def _process_pdf_with_pdfplumber(self, content: bytes, filename: str) -> Dict[str, Any]:
        """Process PDF using pdfplumber"""
        
        def extract_with_plumber():
            with pdfplumber.open(BytesIO(content)) as pdf:
                text_content = []
                page_metadata = []
                tables = []
                
                for i, page in enumerate(pdf.pages):
                    # Extract text
                    page_text = page.extract_text() or ""
                    text_content.append(page_text)
                    
                    # Extract tables from this page
                    page_tables = page.extract_tables()
                    if page_tables:
                        for table in page_tables:
                            if table:
                                tables.append({
                                    'page': i + 1,
                                    'data': table,
                                    'rows': len(table),
                                    'cols': len(table[0]) if table else 0
                                })
                    
                    page_metadata.append({
                        'page_number': i + 1,
                        'char_count': len(page_text),
                        'word_count': len(page_text.split()),
                        'tables_count': len(page_tables) if page_tables else 0
                    })
                
                # Add table text to content
                full_text = '\n\n'.join(text_content)
                if tables:
                    table_text = "\n\nExtracted Tables:\n"
                    for table in tables:
                        table_text += f"\nTable from page {table['page']}:\n"
                        for row in table['data']:
                            table_text += ' | '.join(str(cell) if cell else '' for cell in row) + '\n'
                    full_text += table_text
                
                return {
                    'text': full_text,
                    'page_count': len(pdf.pages),
                    'metadata': {
                        'document_type': 'pdf',
                        'processor': 'pdfplumber',
                        'pages': page_metadata,
                        'total_pages': len(pdf.pages),
                        'tables_extracted': len(tables)
                    }
                }
        
        return await asyncio.to_thread(extract_with_plumber)
    
    async def _process_pdf_with_pypdf2(self, content: bytes, filename: str) -> Dict[str, Any]:
        """Process PDF using PyPDF2"""
        
        def extract_with_pypdf2():
            try:
                pdf_reader = PdfReader(BytesIO(content))
                pages = pdf_reader.pages
                num_pages = len(pages)
            except Exception:
                # Try older PyPDF2 API
                import PyPDF2
                pdf_reader = PyPDF2.PdfFileReader(BytesIO(content))
                num_pages = pdf_reader.numPages
                pages = [pdf_reader.getPage(i) for i in range(num_pages)]
            
            text_content = []
            page_metadata = []
            
            for i, page in enumerate(pages):
                try:
                    page_text = page.extract_text() if hasattr(page, 'extract_text') else page.extractText()
                    text_content.append(page_text)
                    
                    page_metadata.append({
                        'page_number': i + 1,
                        'char_count': len(page_text),
                        'word_count': len(page_text.split())
                    })
                except Exception as e:
                    self.logger.warning(f"Error extracting text from page {i+1}: {e}")
                    text_content.append("")
                    page_metadata.append({'page_number': i + 1, 'char_count': 0, 'word_count': 0})
            
            full_text = '\n\n'.join(text_content)
            
            metadata_info = {}
            try:
                metadata_info = pdf_reader.metadata or {}
                if hasattr(pdf_reader, 'documentInfo'):
                    metadata_info = pdf_reader.documentInfo or {}
            except Exception:
                pass
            
            return {
                'text': full_text,
                'page_count': len(pages),
                'metadata': {
                    'document_type': 'pdf',
                    'processor': 'PyPDF2',
                    'pages': page_metadata,
                    'total_pages': len(pages),
                    'pdf_info': metadata_info
                }
            }
        
        return await asyncio.to_thread(extract_with_pypdf2)
    
    async def process_pdf_fallback(self, content: bytes, filename: str) -> Dict[str, Any]:
        """Fallback PDF processing when libraries aren't available"""
        
        # Extract basic information
        page_count = 1
        
        # Try to estimate page count from PDF structure
        content_str = content.decode('latin-1', errors='ignore')
        page_markers = content_str.count('/Type /Page')
        if page_markers > 0:
            page_count = page_markers
        
        return {
            'text': f"PDF processing libraries not available. Document: {filename} ({page_count} estimated pages)\n\nPlease install PyPDF2 or pdfplumber for full PDF text extraction.",
            'page_count': page_count,
            'metadata': {
                'document_type': 'pdf',
                'processor': 'fallback',
                'note': 'Limited processing - install PDF libraries for full extraction'
            }
        }
    
    async def process_docx(self, content: bytes, filename: str) -> Dict[str, Any]:
        """Process DOCX document"""
        
        def extract_docx():
            doc = DocxDocument(BytesIO(content))
            
            paragraphs = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    paragraphs.append(paragraph.text)
            
            full_text = '\n\n'.join(paragraphs)
            
            # Extract tables
            table_content = []
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells]
                    table_text.append(' | '.join(row_text))
                if table_text:
                    table_content.append('\n'.join(table_text))
            
            if table_content:
                full_text += '\n\nTables:\n' + '\n\n'.join(table_content)
            
            # Extract document properties
            core_props = {}
            try:
                core_props = {
                    'title': doc.core_properties.title or '',
                    'author': doc.core_properties.author or '',
                    'subject': doc.core_properties.subject or '',
                    'created': str(doc.core_properties.created) if doc.core_properties.created else None,
                    'modified': str(doc.core_properties.modified) if doc.core_properties.modified else None
                }
            except Exception:
                pass
            
            return {
                'text': full_text,
                'page_count': max(1, len(paragraphs) // 30),  # Estimate pages
                'metadata': {
                    'document_type': 'docx',
                    'processor': 'python-docx',
                    'paragraph_count': len(paragraphs),
                    'table_count': len(doc.tables),
                    'core_properties': core_props
                }
            }
        
        return await asyncio.to_thread(extract_docx)
    
    async def process_docx_fallback(self, content: bytes, filename: str) -> Dict[str, Any]:
        """Fallback DOCX processing when python-docx isn't available"""
        
        # DOCX is essentially a ZIP file with XML content
        try:
            import zipfile
            import xml.etree.ElementTree as ET
            
            def extract_docx_fallback():
                text_content = []
                
                with zipfile.ZipFile(BytesIO(content)) as docx_zip:
                    # Extract text from document.xml
                    try:
                        with docx_zip.open('word/document.xml') as doc_xml:
                            tree = ET.parse(doc_xml)
                            root = tree.getroot()
                            
                            # Extract text from all text nodes
                            for text_elem in root.iter():
                                if text_elem.text:
                                    text_content.append(text_elem.text)
                    except Exception:
                        pass
                
                full_text = ' '.join(text_content) if text_content else f"DOCX file: {filename} (python-docx required for full extraction)"
                
                return {
                    'text': full_text,
                    'page_count': 1,
                    'metadata': {
                        'document_type': 'docx',
                        'processor': 'fallback_xml',
                        'note': 'Limited processing - install python-docx for full extraction'
                    }
                }
            
            return await asyncio.to_thread(extract_docx_fallback)
            
        except Exception as e:
            return {
                'text': f"DOCX processing libraries not available. Document: {filename}\n\nPlease install python-docx for full DOCX text extraction.",
                'page_count': 1,
                'metadata': {
                    'document_type': 'docx',
                    'processor': 'fallback',
                    'note': 'No processing libraries available'
                }
            }
    
    async def process_txt(self, content: bytes, filename: str) -> Dict[str, Any]:
        """Process TXT document"""
        try:
            # Try different encodings
            text = None
            encoding = None
            
            for enc in ['utf-8', 'utf-16', 'cp1252', 'iso-8859-1', 'ascii']:
                try:
                    text = content.decode(enc)
                    encoding = enc
                    break
                except UnicodeDecodeError:
                    continue
            
            if text is None:
                # Final fallback with error replacement
                text = content.decode('utf-8', errors='replace')
                encoding = 'utf-8-with-errors'
            
            lines = text.split('\n')
            
            # Detect if it might be a different format
            detected_format = self._detect_text_format(text)
            
            return {
                'text': text,
                'page_count': max(1, len(lines) // 50),  # Estimate pages (50 lines per page)
                'metadata': {
                    'document_type': detected_format or 'txt',
                    'processor': 'text_decoder',
                    'line_count': len(lines),
                    'encoding': encoding,
                    'estimated_pages': max(1, len(lines) // 50),
                    'character_count': len(text)
                }
            }
            
        except Exception as e:
            raise DocumentProcessingError(f"TXT processing failed: {str(e)}")
    
    def _detect_text_format(self, text: str) -> Optional[str]:
        """Detect if text file is actually a different format"""
        
        # Check for markdown
        if re.search(r'^#+ ', text, re.MULTILINE) or '```' in text:
            return 'markdown'
        
        # Check for CSV
        lines = text.split('\n')[:5]  # Check first 5 lines
        if all(',' in line for line in lines if line.strip()):
            return 'csv'
        
        # Check for JSON
        text_stripped = text.strip()
        if (text_stripped.startswith('{') and text_stripped.endswith('}')) or \
           (text_stripped.startswith('[') and text_stripped.endswith(']')):
            try:
                import json
                json.loads(text_stripped)
                return 'json'
            except Exception:
                pass
        
        return None
    
    async def process_doc(self, content: bytes, filename: str) -> Dict[str, Any]:
        """Process DOC document (legacy Word format)"""
        
        # DOC format is complex and requires specialized libraries
        # For production, consider using libraries like python-docx2txt or antiword
        
        try:
            # Try to extract any readable text
            text_content = content.decode('latin-1', errors='ignore')
            
            # Remove binary characters and keep only readable text
            readable_text = re.sub(r'[^\x20-\x7E\n\r\t]', ' ', text_content)
            readable_text = re.sub(r'\s+', ' ', readable_text).strip()
            
            # If we got some reasonable text
            if len(readable_text) > 100:
                return {
                    'text': readable_text,
                    'page_count': 1,
                    'metadata': {
                        'document_type': 'doc',
                        'processor': 'binary_extraction',
                        'note': 'Basic text extraction - results may be incomplete. Consider converting to DOCX.'
                    }
                }
        except Exception:
            pass
        
        return {
            'text': f"Legacy DOC format detected: {filename}\n\nFor best results, please convert to DOCX or PDF format. DOC processing requires specialized libraries like python-docx2txt.",
            'page_count': 1,
            'metadata': {
                'document_type': 'doc',
                'processor': 'fallback',
                'processing_note': 'Legacy DOC format - conversion recommended'
            }
        }
    
    async def process_rtf(self, content: bytes, filename: str) -> Dict[str, Any]:
        """Process RTF document"""
        try:
            # RTF is a plain text format with formatting codes
            text = content.decode('utf-8', errors='ignore')
            
            # Basic RTF cleaning - remove RTF control codes
            # This is a simplified approach
            cleaned_text = re.sub(r'\$$a-z]+\d*\s?', '', text)  # Remove RTF commands
            cleaned_text = re.sub(r'[{}]', '', cleaned_text)     # Remove braces
            cleaned_text = re.sub(r'\s+', ' ', cleaned_text)     # Normalize whitespace
            cleaned_text = cleaned_text.strip()
            
            return {
                'text': cleaned_text,
                'page_count': max(1, len(cleaned_text) // 2500),
                'metadata': {
                    'document_type': 'rtf',
                    'processor': 'basic_rtf_cleaner',
                    'processing_note': 'Basic RTF text extraction - specialized RTF libraries recommended for better results'
                }
            }
            
        except Exception as e:
            raise DocumentProcessingError(f"RTF processing failed: {str(e)}")
    
    async def process_unknown(self, content: bytes, filename: str) -> Dict[str, Any]:
        """Process unknown/generic file types"""
        
        # Try to extract any readable text
        try:
            # Try UTF-8 first
            text = content.decode('utf-8', errors='ignore')
        except Exception:
            # Fallback to latin-1
            text = content.decode('latin-1', errors='ignore')
        
        # Clean up binary characters
        readable_text = re.sub(r'[^\x20-\x7E\n\r\t]', ' ', text)
        readable_text = re.sub(r'\s+', ' ', readable_text).strip()
        
        if len(readable_text) < 50:
            readable_text = f"Unknown file format: {filename}\n\nFile appears to be binary or uses an unsupported format. Please convert to PDF, DOCX, or TXT."
        
        return {
            'text': readable_text,
            'page_count': 1,
            'metadata': {
                'document_type': 'unknown',
                'processor': 'generic_text_extraction',
                'original_filename': filename,
                'file_size': len(content),
                'processing_note': 'Generic text extraction from unknown format'
            }
        }
    
    def _get_processor_by_extension(self, filename: str) -> Optional[callable]:
        """Get processor based on file extension"""
        
        extension = Path(filename).suffix.lower()
        
        extension_map = {
            '.pdf': self.process_pdf if PYPDF2_AVAILABLE or PDFPLUMBER_AVAILABLE else self.process_pdf_fallback,
            '.docx': self.process_docx if DOCX_AVAILABLE else self.process_docx_fallback,
            '.doc': self.process_doc,
            '.txt': self.process_txt,
            '.rtf': self.process_rtf,
            '.md': self.process_txt,
            '.markdown': self.process_txt,
            '.csv': self.process_txt,
            '.json': self.process_txt
        }
        
        return extension_map.get(extension)
    
    def _detect_mime_type(self, content: bytes, filename: str) -> str:
        """Detect MIME type of document"""
        
        if MAGIC_AVAILABLE:
            try:
                mime_type = magic.from_buffer(content, mime=True)
                return mime_type
            except Exception as e:
                self.logger.warning(f"python-magic failed: {e}")
        
        # Fallback to filename extension and magic bytes
        extension = Path(filename).suffix.lower()
        
        # Check magic bytes first
        if content.startswith(b'%PDF'):
            return 'application/pdf'
        elif content.startswith(b'PK\x03\x04') and extension == '.docx':
            return 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        elif content.startswith(b'\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1'):
            return 'application/msword'
        elif content.startswith(b'{\\rtf'):
            return 'text/rtf'
        
        # Fallback to mimetypes based on extension
        mime_type, _ = mimetypes.guess_type(filename)
        return mime_type or 'application/octet-stream'
    
    def _generate_document_id(self) -> str:
        """Generate unique document ID"""
        return f"doc_{uuid.uuid4().hex[:12]}"
    
    def _calculate_file_hash(self, content: bytes) -> str:
        """Calculate SHA-256 hash of file content"""
        return hashlib.sha256(content).hexdigest()
    
    def get_supported_formats(self) -> List[str]:
        """Get list of supported formats"""
        formats = ['txt', 'rtf', 'md']
        
        if PYPDF2_AVAILABLE or PDFPLUMBER_AVAILABLE:
            formats.append('pdf')
        
        if DOCX_AVAILABLE:
            formats.append('docx')
        
        formats.extend(['doc', 'csv', 'json'])  # These have basic support
        
        return formats
    
    def get_processor_info(self) -> Dict[str, Any]:
        """Get information about available processors"""
        return {
            'pdf_processors': {
                'pdfplumber': PDFPLUMBER_AVAILABLE,
                'PyPDF2': PYPDF2_AVAILABLE,
                'fallback': True
            },
            'docx_processors': {
                'python-docx': DOCX_AVAILABLE,
                'xml_fallback': True,
                'basic_fallback': True
            },
            'other_formats': {
                'txt': True,
                'rtf': True,
                'doc': True,
                'unknown': True
            },
            'utilities': {
                'python-magic': MAGIC_AVAILABLE,
                'aiofiles': AIOFILES_AVAILABLE,
                'intelligent_chunker': CHUNKER_AVAILABLE
            },
            'supported_formats': self.get_supported_formats()
        }

class DocumentTypeDetector:
    """
    Advanced document type detection and validation with fallbacks
    """
    
    @staticmethod
    def detect_document_type(content: bytes, filename: str) -> str:
        """Detect document type with high accuracy"""
        
        # Check magic bytes
        magic_signatures = {
            b'%PDF': 'pdf',
            b'PK\x03\x04': 'office',  # ZIP-based format (could be docx, xlsx, etc.)
            b'\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1': 'ole',  # OLE format (doc, xls, ppt)
            b'{\\rtf': 'rtf',
            b'\x7fELF': 'binary',
            b'\x89PNG': 'image',
            b'\xff\xd8\xff': 'image'
        }
        
        for signature, doc_type in magic_signatures.items():
            if content.startswith(signature):
                # For ZIP-based formats, check filename extension
                if doc_type == 'office':
                    extension = Path(filename).suffix.lower()
                    if extension == '.docx':
                        return 'docx'
                    elif extension == '.xlsx':
                        return 'xlsx'
                    elif extension == '.pptx':
                        return 'pptx'
                    else:
                        return 'zip_archive'
                
                # For OLE formats, check filename extension
                elif doc_type == 'ole':
                    extension = Path(filename).suffix.lower()
                    if extension == '.doc':
                        return 'doc'
                    elif extension == '.xls':
                        return 'xls'
                    elif extension == '.ppt':
                        return 'ppt'
                    else:
                        return 'ole_document'
                
                return doc_type
        
        # Check if it's readable text
        try:
            text_sample = content[:1000].decode('utf-8')
            # Check for specific text formats
            if text_sample.strip().startswith('#') or '```' in text_sample:
                return 'markdown'
            elif text_sample.count(',') > text_sample.count(' ') / 2:
                return 'csv'
            elif text_sample.strip().startswith(('{', '[')):
                return 'json'
            else:
                return 'txt'
        except UnicodeDecodeError:
            pass
        
        # Fallback to extension
        extension = Path(filename).suffix.lower()
        extension_map = {
            '.pdf': 'pdf',
            '.docx': 'docx',
            '.doc': 'doc',
            '.txt': 'txt',
            '.rtf': 'rtf',
            '.md': 'markdown',
            '.csv': 'csv',
            '.json': 'json'
        }
        
        return extension_map.get(extension, 'unknown')
    
    @staticmethod
    def validate_document(content: bytes, expected_type: str) -> Tuple[bool, str]:
        """
        Validate that document matches expected type
        
        Returns:
            (is_valid, detected_type)
        """
        detected_type = DocumentTypeDetector.detect_document_type(content, f"file.{expected_type}")
        is_valid = detected_type == expected_type or (
            expected_type == 'txt' and detected_type in ['txt', 'markdown', 'csv', 'json']
        )
        
        return is_valid, detected_type
    
    @staticmethod
    def get_file_info(content: bytes, filename: str) -> Dict[str, Any]:
        """Get comprehensive file information"""
        
        detected_type = DocumentTypeDetector.detect_document_type(content, filename)
        is_valid, _ = DocumentTypeDetector.validate_document(content, detected_type)
        
        # Try to determine if it's readable text
        is_text = False
        encoding = None
        try:
            content[:1000].decode('utf-8')
            is_text = True
            encoding = 'utf-8'
        except UnicodeDecodeError:
            try:
                content[:1000].decode('latin-1')
                is_text = True
                encoding = 'latin-1'
            except Exception:
                pass
        
        return {
            'filename': filename,
            'detected_type': detected_type,
            'file_size': len(content),
            'is_valid': is_valid,
            'is_text': is_text,
            'encoding': encoding,
            'magic_bytes': content[:16].hex() if len(content) >= 16 else content.hex(),
            'estimated_pages': max(1, len(content) // 2500) if is_text else 'unknown'
        }

# Utility functions
def create_mock_processor() -> MultiDocumentProcessor:
    """Create processor with all fallback implementations"""
    return MultiDocumentProcessor()

async def process_file_simple(file_content: bytes, filename: str) -> Dict[str, Any]:
    """Simple file processing function for basic use cases"""
    
    processor = MultiDocumentProcessor()
    try:
        document, chunks = await processor.process_document(
            file_content, filename, "default_session"
        )
        
        return {
            'success': True,
            'document': document.to_dict() if hasattr(document, 'to_dict') else str(document),
            'chunks': [chunk.to_dict() if hasattr(chunk, 'to_dict') else str(chunk) for chunk in chunks],
            'text': document.content,
            'word_count': document.word_count,
            'page_count': document.page_count
        }
    except Exception as e:
        return {
            'success': False,
            'error': str(e),
            'document': None,
            'chunks': [],
            'text': '',
            'word_count': 0,
            'page_count': 0
        }
